"""Document chat service — upload, analyze, and chat with legal documents."""
from __future__ import annotations

import asyncio
import io
import json
import logging
import uuid
from typing import Optional

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.core import settings
from app.models.doc_sessions import DocMessage, DocSession
from app.services.document_service import _get_genai_client, _strip_json_fences

logger = logging.getLogger(__name__)


# ── Text Extraction ──────────────────────────────────────────────


async def extract_text_from_bytes(content: bytes, content_type: str, file_name: str) -> str:
    """Extract text from uploaded file bytes."""
    ct = (content_type or "").lower()

    if ct in ("text/plain", "text/markdown", "text/csv"):
        return content.decode("utf-8", errors="replace")

    if ct == "application/pdf" or file_name.lower().endswith(".pdf"):
        return await _extract_pdf_text(content)

    if ct in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ) or file_name.lower().endswith((".docx", ".doc")):
        return await _extract_docx_text(content)

    # Fallback: try as plain text
    try:
        return content.decode("utf-8", errors="replace")
    except Exception:
        return ""


async def _extract_pdf_text(content: bytes) -> str:
    """Extract text from PDF using Gemini's vision (handles scanned docs too)."""
    if not settings.gemini_api_key:
        return "[PDF extraction requires GEMINI_API_KEY]"

    import base64

    client = _get_genai_client()
    b64 = base64.standard_b64encode(content).decode()

    response = await client.aio.models.generate_content(
        model=settings.gemini_lite_model,
        contents=[
            {
                "parts": [
                    {"inline_data": {"mime_type": "application/pdf", "data": b64}},
                    {"text": "Extract ALL text from this document exactly as written. Preserve paragraph structure. Do not summarize or interpret — just extract the raw text content."},
                ]
            }
        ],
    )
    return response.text.strip()


async def _extract_docx_text(content: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    import docx as python_docx

    doc = python_docx.Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n\n".join(paragraphs)


# ── Analysis ─────────────────────────────────────────────────────


async def analyze_document(text: str) -> dict:
    """Deep analysis of a legal document using Gemini."""
    if not settings.gemini_api_key or not text.strip():
        return {"error": "Cannot analyze: no API key or empty text"}

    client = _get_genai_client()

    prompt = f"""You are an expert Indian legal analyst. Analyze this legal document thoroughly.

Document text:
{text[:12000]}

Provide a JSON response with:
{{
  "document_type": "type of document (e.g., rental agreement, NDA, contract, legal notice, power of attorney, court order, etc.)",
  "summary": "2-3 paragraph summary of the document",
  "parties": [{{"name": "...", "role": "..."}}],
  "key_terms": [{{"label": "...", "value": "...", "clause_ref": "..."}}],
  "effective_date": "date or null",
  "termination_date": "date or null",
  "obligations": [{{"party": "...", "obligation": "..."}}],
  "risks": ["list of potential risks or unfavorable clauses"],
  "important_clauses": [{{"title": "...", "summary": "..."}}],
  "governing_law": "applicable law/jurisdiction",
  "recommendations": ["actionable recommendations for the reader"]
}}

Respond ONLY with valid JSON, no markdown fences."""

    response = await client.aio.models.generate_content(
        model=settings.gemini_model,
        contents=prompt,
    )

    raw = _strip_json_fences(response.text)
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        return {"summary": response.text.strip(), "error": "Could not parse structured analysis"}


# ── Session Management ───────────────────────────────────────────


async def create_session(
    session: AsyncSession,
    file_name: str,
    content: bytes,
    content_type: str,
    gcs_object_key: Optional[str] = None,
    org_id: Optional[str] = None,
    user_id: Optional[str] = None,
) -> DocSession:
    """Create a doc session and kick off extraction + analysis in background."""
    doc_session = DocSession(
        file_name=file_name,
        content_type=content_type,
        byte_size=len(content),
        gcs_object_key=gcs_object_key,
        org_id=org_id,
        user_id=user_id,
        status="processing",
    )
    session.add(doc_session)
    await session.commit()
    await session.refresh(doc_session)

    return doc_session


async def process_document_inline(
    session: AsyncSession,
    doc_session: DocSession,
    content: bytes,
    content_type: str,
    file_name: str,
) -> None:
    """Extract text and analyze the document inline (within the request)."""
    # Extract text
    try:
        text = await extract_text_from_bytes(content, content_type, file_name)
        doc_session.extracted_text = text
    except Exception as e:
        logger.exception(f"[DocChat] Text extraction failed for {doc_session.id}")
        doc_session.status = "failed"
        doc_session.error = f"Text extraction failed: {e}"
        await session.commit()
        return

    # Analyze
    try:
        analysis = await analyze_document(text)
        doc_session.analysis = analysis
        doc_session.status = "ready"
    except Exception as e:
        logger.exception(f"[DocChat] Analysis failed for {doc_session.id}")
        doc_session.status = "ready"  # Still usable for chat, just no analysis
        doc_session.analysis = {"error": str(e)}

    await session.commit()
    logger.info(f"[DocChat] Inline processing complete for {doc_session.id}: {doc_session.status}")


async def _process_document(
    session_id: str,
    content: bytes,
    content_type: str,
    file_name: str,
) -> None:
    """Background task: extract text and analyze the document."""
    from app.db.session import async_session_factory

    async with async_session_factory() as session:
        doc_session = await session.get(DocSession, session_id)
        if not doc_session:
            logger.error(f"[DocChat] Session {session_id} not found for processing")
            return

        # Extract text
        try:
            text = await extract_text_from_bytes(content, content_type, file_name)
            doc_session.extracted_text = text
        except Exception as e:
            logger.exception(f"[DocChat] Text extraction failed for {session_id}")
            doc_session.status = "failed"
            doc_session.error = f"Text extraction failed: {e}"
            await session.commit()
            return

        # Analyze
        try:
            analysis = await analyze_document(text)
            doc_session.analysis = analysis
            doc_session.status = "ready"
        except Exception as e:
            logger.exception(f"[DocChat] Analysis failed for {session_id}")
            doc_session.status = "ready"  # Still usable for chat, just no analysis
            doc_session.analysis = {"error": str(e)}

        await session.commit()
        logger.info(f"[DocChat] Processing complete for {session_id}: {doc_session.status}")


async def get_session(session: AsyncSession, session_id: str) -> Optional[DocSession]:
    result = await session.execute(
        select(DocSession)
        .options(selectinload(DocSession.messages))
        .where(DocSession.id == session_id)
    )
    return result.scalar_one_or_none()


async def list_sessions(
    session: AsyncSession,
    limit: int = 20,
    org_id: Optional[str] = None,
    user_id: Optional[str] = None,
) -> list[DocSession]:
    stmt = select(DocSession)
    if org_id:
        stmt = stmt.where(DocSession.org_id == org_id)
    if user_id:
        stmt = stmt.where(DocSession.user_id == user_id)
    result = await session.execute(
        stmt.order_by(DocSession.created_at.desc()).limit(limit)
    )
    return list(result.scalars().all())


# ── Chat ─────────────────────────────────────────────────────────


async def chat(
    session: AsyncSession,
    doc_session: DocSession,
    user_message: str,
) -> tuple[str, list[dict]]:
    """Chat with a document — uses document text + chat history as context.

    Returns (assistant_text, citations) where citations is a list of
    ``{"text": "...", "clause_ref": "..."}`` dicts extracted from the document.
    """
    if not settings.gemini_api_key:
        return "Chat requires GEMINI_API_KEY to be configured.", []

    # Save user message
    user_msg = DocMessage(
        session_id=doc_session.id,
        role="user",
        content=user_message,
    )
    session.add(user_msg)
    await session.commit()

    # Build context
    doc_text = (doc_session.extracted_text or "")[:10000]
    analysis_json = json.dumps(doc_session.analysis or {}, indent=2)[:3000]

    # Build chat history (last 10 messages)
    history = doc_session.messages[-10:] if doc_session.messages else []
    history_text = "\n".join(
        f"{'User' if m.role == 'user' else 'Assistant'}: {m.content}"
        for m in history
    )

    client = _get_genai_client()

    prompt = f"""You are an expert Indian legal assistant. A user has uploaded a document and is asking questions about it.

DOCUMENT ({doc_session.file_name}):
{doc_text}

ANALYSIS:
{analysis_json}

CONVERSATION HISTORY:
{history_text}

User's question: {user_message}

Instructions:
- Answer based ONLY on the document content. If the answer isn't in the document, say so.
- Be specific — cite clause numbers, sections, or exact phrases from the document.
- If the user asks about risks or recommendations, be thorough and practical.
- Use simple language but be legally precise.
- If relevant, mention applicable Indian laws or precedents.

IMPORTANT: Respond with a JSON object in this exact format (no markdown fences):
{{
  "answer": "Your full answer in markdown format",
  "citations": [
    {{
      "text": "Exact quoted text from the document that supports this part of your answer",
      "clause_ref": "Section/clause reference if identifiable, or null"
    }}
  ]
}}

Rules for citations:
- Each citation's "text" must be a VERBATIM excerpt from the document (not paraphrased).
- Include 2-5 citations for substantive answers. Use fewer only if the document lacks relevant content.
- Keep each citation excerpt to 1-3 sentences — enough for context but not entire paragraphs.
- If no specific text supports the answer, use an empty citations array.
- The "clause_ref" should reference the section, clause, or article number if one is apparent."""

    response = await client.aio.models.generate_content(
        model=settings.gemini_model,
        contents=prompt,
    )

    raw = _strip_json_fences(response.text.strip())

    # Parse structured response with graceful fallback
    citations: list[dict] = []
    try:
        parsed = json.loads(raw)
        assistant_text = parsed.get("answer", raw)
        citations = [
            {"text": c.get("text", ""), "clause_ref": c.get("clause_ref")}
            for c in parsed.get("citations", [])
            if isinstance(c, dict) and c.get("text")
        ]
    except (json.JSONDecodeError, AttributeError):
        assistant_text = response.text.strip()

    # Save assistant response with citations in metadata
    assistant_msg = DocMessage(
        session_id=doc_session.id,
        role="assistant",
        content=assistant_text,
        extra_data={"citations": citations} if citations else None,
    )
    session.add(assistant_msg)
    await session.commit()

    return assistant_text, citations
